# -*- coding:utf-8 -*-
#txt转doc

import docx
import os
import sys
import re


class CTxt2Docx(object):
    def __init__(self):
        self.m_doc = docx.Document()

    def write(self, txt_path, save_path):
        if os.path.exists(txt_path) is False:
            raise SystemExit("[Error] txtpath is not exist")
        content = ""

        def read_callback(line):
            nonlocal content
            line = re.sub(r"{{(.*)?}}", "", line)
            content += line

        self.read(txt_path, read_callback)
        self.m_doc.add_paragraph(content)
        self.m_doc.save(save_path)

    def write2(self, txt_path, save_path):
        if os.path.exists(txt_path) is False:
            raise SystemExit("[Error] txtpath is not exist")
        content = ""

        def read_callback(line):
            nonlocal content
            search = re.search(r"{{(.*?):(.*?)}}", line)
            if search is not None:
                line = re.sub(r"{{(.*)?}}", "", line)
                p = self.m_doc.add_paragraph()
                run = p.add_run(r"你好")
                color = search.groups()[1]
                run.font.color.rgb = self.change_color(color)
            else:
                content += line

        self.read(txt_path, read_callback)
        self.m_doc.add_paragraph(content)
        self.m_doc.save(save_path)

    def change_color(self, src):
        if src == "red":
            return docx.shared.RGBColor(255, 0, 0)
        elif src == "blue":
            return docx.shared.RGBColor(0, 0, 255)
        elif src == "green":
            return docx.shared.RGBColor(0, 255, 0)
        else:
            return docx.shared.RGBColor(0, 0, 0)

    def read(self, txt_path, callback):
        # read txt
        fp = open(txt_path, encoding="utf8")
        line = fp.readline()
        while line:
            callback(line)
            line = fp.readline()


if __name__ == "__main__":
    writer = CTxt2Docx()
    writer.write("../todo.txt", "todo.docx")
